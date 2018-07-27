import docx
import os


def word_read():
    file = docx.Document("5500.docx")
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段
    i = 0
    # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)
    #     i += 1
    # print(i)
    # 输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        print("第" + str(i) + "e：" + file.paragraphs[i].text)


if __name__ == '__main__':
    word_read()
